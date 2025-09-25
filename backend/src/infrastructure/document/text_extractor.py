"""
Text Extraction Service for Multi-Format Document Processing
Supports PDF, DOCX, TXT files with security validation
"""

import io
import magic
from pathlib import Path
from typing import Union, Dict, Any
import logging
from PyPDF2 import PdfReader
from docx import Document
import re

from ...shared.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Secure text extraction from multiple file formats
    """
    
    # Supported MIME types
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'text/html': 'html'
    }
    
    # Maximum file size (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    def __init__(self):
        self.mime_detector = magic.Magic(mime=True)
    
    def extract_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from uploaded file with validation
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Dict containing extracted text and metadata
            
        Raises:
            DocumentProcessingError: If file is invalid or processing fails
        """
        try:
            # Validate file size
            if len(file_content) > self.MAX_FILE_SIZE:
                raise DocumentProcessingError(
                    f"File size {len(file_content)} exceeds maximum {self.MAX_FILE_SIZE} bytes"
                )
            
            # Detect MIME type
            mime_type = self.mime_detector.from_buffer(file_content)
            logger.info(f"Detected MIME type: {mime_type} for file: {filename}")
            
            # Validate MIME type
            if mime_type not in self.SUPPORTED_MIME_TYPES:
                raise DocumentProcessingError(
                    f"Unsupported file type: {mime_type}. Supported types: {list(self.SUPPORTED_MIME_TYPES.keys())}"
                )
            
            # Extract text based on file type
            file_type = self.SUPPORTED_MIME_TYPES[mime_type]
            
            if file_type == 'pdf':
                text = self._extract_from_pdf(file_content)
            elif file_type == 'docx':
                text = self._extract_from_docx(file_content)
            elif file_type == 'txt':
                text = self._extract_from_txt(file_content)
            elif file_type == 'html':
                text = self._extract_from_html(file_content)
            else:
                raise DocumentProcessingError(f"No extractor implemented for {file_type}")
            
            # Clean and validate extracted text
            cleaned_text = self._clean_text(text)
            
            if not cleaned_text.strip():
                raise DocumentProcessingError("No text content found in document")
            
            # Return extraction results
            result = {
                'text': cleaned_text,
                'word_count': len(cleaned_text.split()),
                'character_count': len(cleaned_text),
                'file_type': file_type,
                'mime_type': mime_type,
                'filename': filename,
                'file_size': len(file_content)
            }
            
            logger.info(f"Successfully extracted {result['word_count']} words from {filename}")
            return result
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            if isinstance(e, DocumentProcessingError):
                raise
            raise DocumentProcessingError(f"Text extraction failed: {str(e)}")
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue
            
            if not text_parts:
                raise DocumentProcessingError("No readable text found in PDF")
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise DocumentProcessingError(f"PDF extraction failed: {str(e)}")
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            document = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if not text_parts:
                raise DocumentProcessingError("No readable text found in DOCX")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise DocumentProcessingError(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise DocumentProcessingError(f"TXT extraction failed: {str(e)}")
    
    def _extract_from_html(self, file_content: bytes) -> str:
        """Extract text from HTML file (basic implementation)"""
        try:
            html_text = file_content.decode('utf-8', errors='ignore')
            
            # Remove HTML tags (basic regex approach)
            text = re.sub(r'<[^>]+>', '', html_text)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text
            
        except Exception as e:
            raise DocumentProcessingError(f"HTML extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\"\']+', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def validate_file(self, file_content: bytes, filename: str) -> bool:
        """
        Validate file without extracting text
        
        Returns:
            bool: True if file is valid for processing
        """
        try:
            # Check file size
            if len(file_content) > self.MAX_FILE_SIZE:
                return False
            
            # Check MIME type
            mime_type = self.mime_detector.from_buffer(file_content)
            
            return mime_type in self.SUPPORTED_MIME_TYPES
            
        except Exception:
            return False


# Global text extractor instance
text_extractor = TextExtractor()